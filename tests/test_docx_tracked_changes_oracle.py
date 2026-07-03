"""Kill-proof tests for the ``docx_tracked_changes_readback`` oracle.

These tests exercise ``kairo.oracles.verify_docx_tracked_changes`` against a REAL
``.docx`` produced by the REAL tracked-changes engine
(``sidecar/parsers/adeu_bridge.py::_python_docx_tracked_fallback``), which emits
genuine OOXML ``w:ins``/``w:del`` revision markup. No mocks.

Each "killproof" test intentionally breaks the produced document (drops a
revision, strips an author/date, or tampers the recovered text) and asserts the
oracle FAILS — proving the oracle is load-bearing (specs/VERIFICATION_ORACLES.md
"an oracle that cannot be shown to fail on bad input is itself rigged").
"""
from __future__ import annotations

import importlib.util
import os
import shutil

import pytest
from docx import Document
from docx.oxml.ns import qn

from kairo.oracles.docx_tracked_changes import (
    extract_revisions,
    reconstruct_original_and_final,
    verify_docx_tracked_changes,
)

# --- Load the REAL tracked-changes engine standalone (avoid heavy package init) ---
_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_ENGINE_PATH = os.path.join(
    _REPO_ROOT, "kairo-sidecar", "sidecar", "parsers", "adeu_bridge.py"
)
_spec = importlib.util.spec_from_file_location("_kairo_adeu_bridge", _ENGINE_PATH)
_engine = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_engine)
_apply_tracked = _engine._python_docx_tracked_fallback  # real engine, real w:ins/w:del

AUTHOR = "Kairo Legal"
_ORIGINAL = (
    "This Agreement is governed by the laws of Delaware. "
    "Either party may terminate with 90 days notice. "
    "Liability is unlimited for all claims."
)
_FINAL = (
    "This Agreement is governed by the laws of the State of New York. "
    "Either party may terminate with 30 days notice. "
    "Liability is capped at fees paid in the prior 12 months."
)
_EDITS = [
    {"target_text": "the laws of Delaware", "new_text": "the laws of the State of New York"},
    {"target_text": "90 days", "new_text": "30 days"},
    {
        "target_text": "Liability is unlimited for all claims.",
        "new_text": "Liability is capped at fees paid in the prior 12 months.",
    },
]
_EXPECTED = [
    {"old": "the laws of Delaware", "new": "the laws of the State of New York"},
    {"old": "90 days", "new": "30 days"},
    {
        "old": "Liability is unlimited for all claims.",
        "new": "Liability is capped at fees paid in the prior 12 months.",
    },
]


def _make_base(path: str) -> None:
    doc = Document()
    doc.add_paragraph("This Agreement is governed by the laws of Delaware.")
    doc.add_paragraph("Either party may terminate with 90 days notice.")
    doc.add_paragraph("Liability is unlimited for all claims.")
    doc.save(path)


def _redline(tmp_path, name: str, edits) -> str:
    base = os.path.join(str(tmp_path), "base.docx")
    if not os.path.exists(base):
        _make_base(base)
    out = os.path.join(str(tmp_path), name)
    res = _apply_tracked(base, edits, out, AUTHOR)
    assert res["ok"], res
    return out


# --------------------------- positive behaviour ---------------------------

def test_positive_readback(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    assert verify_docx_tracked_changes(
        out,
        _EXPECTED,
        require_author=True,
        require_date=True,
        original_text=_ORIGINAL,
        final_text=_FINAL,
    ) is True


def test_extract_revisions_have_author_and_date(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    revs = extract_revisions(out)
    assert len([r for r in revs if r.kind == "del"]) == 3
    assert len([r for r in revs if r.kind == "ins"]) == 3
    for r in revs:
        assert r.author == AUTHOR
        assert r.date  # non-empty timestamp


def test_reconstruction_roundtrip(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    original, final = reconstruct_original_and_final(out)
    # normalized comparison (oracle normalizes whitespace/NFC internally)
    import re
    import unicodedata

    def n(s):
        return re.sub(r"\s+", " ", unicodedata.normalize("NFC", s)).strip()

    assert original == n(_ORIGINAL)
    assert final == n(_FINAL)


# ------------------------------ kill-proofs -------------------------------

def test_killproof_dropped_revision(tmp_path):
    # Engine applies only 2 of the 3 intended edits => a required revision is absent.
    out = _redline(tmp_path, "partial.docx", _EDITS[:2])
    with pytest.raises(AssertionError, match="not found"):
        verify_docx_tracked_changes(out, _EXPECTED)


def test_killproof_missing_author(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    tampered = os.path.join(str(tmp_path), "noauthor.docx")
    shutil.copy(out, tampered)
    doc = Document(tampered)
    for ins in doc.element.body.iter(qn("w:ins")):
        if ins.get(qn("w:author")) is not None:
            del ins.attrib[qn("w:author")]
            break
    doc.save(tampered)
    with pytest.raises(AssertionError, match="w:author"):
        verify_docx_tracked_changes(tampered, _EXPECTED, require_author=True)


def test_killproof_missing_date(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    tampered = os.path.join(str(tmp_path), "nodate.docx")
    shutil.copy(out, tampered)
    doc = Document(tampered)
    for dele in doc.element.body.iter(qn("w:del")):
        if dele.get(qn("w:date")) is not None:
            del dele.attrib[qn("w:date")]
            break
    doc.save(tampered)
    with pytest.raises(AssertionError, match="w:date"):
        verify_docx_tracked_changes(tampered, _EXPECTED, require_date=True)


def test_killproof_original_recovery_mismatch(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    with pytest.raises(AssertionError, match="recover the original"):
        verify_docx_tracked_changes(out, _EXPECTED, original_text="WRONG ORIGINAL TEXT")


def test_killproof_final_mismatch(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    with pytest.raises(AssertionError, match="expected final text"):
        verify_docx_tracked_changes(out, _EXPECTED, final_text="WRONG FINAL TEXT")


def test_killproof_extra_unintended_revision(tmp_path):
    # Only claim one change is expected, but the document has three.
    out = _redline(tmp_path, "full.docx", _EDITS)
    with pytest.raises(AssertionError, match="unintended tracked revisions"):
        verify_docx_tracked_changes(out, _EXPECTED[:1], forbid_extra_revisions=True)


def test_empty_expected_changes_rejected(tmp_path):
    out = _redline(tmp_path, "full.docx", _EDITS)
    with pytest.raises(AssertionError, match="non-empty list"):
        verify_docx_tracked_changes(out, [])
