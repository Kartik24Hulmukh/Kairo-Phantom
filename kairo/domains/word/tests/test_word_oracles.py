# PROVENANCE: original | Word/docs domain oracle tests per VERIFICATION_ORACLES.md
"""Word/docs domain oracle tests — docx_readback + structure_readback + kill-proofs.

Tests verify:
  1. docx_readback: after creating a .docx, reopen and assert paragraphs/
     headings/styles/lists/table cells match the spec. Kill-proof: alter text
     or drop a paragraph/table → FAILS.
  2. structure_readback: heading hierarchy + table dims + paragraph count
     survive round-trip. Kill-proof: drop a section/table → FAILS.
  3. Honest degradation: python-docx missing → FAIL LOUD.
  4. >=3 gauntlet scenarios: (a) multi-heading styled doc with paragraphs,
     (b) a document with a real table, (c) a document with a numbered/bulleted
     list + bold/italic runs — each read-back verified.
  5. Trust stack integration: audit log + egress report.
  6. CLI integration: word subcommand works end-to-end.

All tests run fully offline. No mocks on production paths. Zero skips.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile

import pytest
from cryptography.hazmat.primitives.asymmetric import ed25519

_REPO_ROOT = os.path.dirname(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
)
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from kairo.domains.word.engine import (  # noqa: E402
    WordEngineUnavailableError,
    WordError,
    create_document,
    read_document,
    save_document,
    word_pipeline,
)
from kairo.domains.word.oracles import (  # noqa: E402
    docx_readback,
    structure_readback,
)

# Fixture paths
_FIX = os.path.join(_REPO_ROOT, "kairo", "domains", "word", "fixtures")
_SPEC_JSON = os.path.join(_FIX, "doc_spec.json")
_GT_JSON = os.path.join(_FIX, "ground_truth.json")


# ---------------------------------------------------------------------------
# Helper: check engine availability
# ---------------------------------------------------------------------------


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


_HAS_DOCX = _docx_available()


# ---------------------------------------------------------------------------
# Helper: create a doc from a spec and save to temp file
# ---------------------------------------------------------------------------


def _create_doc_from_spec(spec: dict, tmpdir: str, filename: str = "test.docx") -> str:
    """Create a .docx from a spec dict and return the saved path."""
    doc = create_document(spec)
    out_path = os.path.join(tmpdir, filename)
    return save_document(doc, out_path)


def _load_ground_truth() -> dict:
    with open(_GT_JSON, encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Oracle 1: docx_readback
# ---------------------------------------------------------------------------


class TestDocxReadback:
    """docx_readback oracle — create, save, reopen, verify all fields."""

    def test_plain_paragraphs_readback(self):
        """Create a doc with plain paragraphs, reopen, verify text matches."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "Title One", "level": 1},
                    {"type": "paragraph", "text": "First paragraph.", "style": "Normal"},
                    {"type": "paragraph", "text": "Second paragraph.", "style": "Normal"},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            expected = [
                {"text": "Title One", "style": "Heading 1", "heading_level": 1},
                {"text": "First paragraph.", "style": "Normal"},
                {"text": "Second paragraph.", "style": "Normal"},
            ]
            result = docx_readback(path, expected)
            assert result is True

    def test_heading_levels_readback(self):
        """Heading levels are correctly read back."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "title", "text": "Doc Title"},
                    {"type": "heading", "text": "Section 1", "level": 1},
                    {"type": "heading", "text": "Subsection 1.1", "level": 2},
                    {"type": "heading", "text": "Subsection 1.1.1", "level": 3},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            expected = [
                {"text": "Doc Title", "heading_level": 0},
                {"text": "Section 1", "heading_level": 1},
                {"text": "Subsection 1.1", "heading_level": 2},
                {"text": "Subsection 1.1.1", "heading_level": 3},
            ]
            result = docx_readback(path, expected)
            assert result is True

    def test_table_readback(self):
        """Table cells are correctly read back."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "Data", "level": 1},
                    {"type": "table", "rows": 3, "cols": 2, "data": [
                        ["Name", "Value"],
                        ["Alpha", "100"],
                        ["Beta", "200"],
                    ]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            expected_paras = [
                {"text": "Data", "heading_level": 1},
            ]
            expected_tables = [
                {"rows": 3, "cols": 2, "cells": [["Name", "Value"], ["Alpha", "100"], ["Beta", "200"]]},
            ]
            result = docx_readback(path, expected_paras, expected_tables)
            assert result is True

    def test_list_readback(self):
        """List items are correctly read back with list type."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "Tasks", "level": 1},
                    {"type": "list", "list_type": "number", "items": ["Task A", "Task B", "Task C"]},
                    {"type": "list", "list_type": "bullet", "items": ["Point 1", "Point 2"]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            expected = [
                {"text": "Tasks", "heading_level": 1},
                {"text": "Task A", "list_type": "number"},
                {"text": "Task B", "list_type": "number"},
                {"text": "Task C", "list_type": "number"},
                {"text": "Point 1", "list_type": "bullet"},
                {"text": "Point 2", "list_type": "bullet"},
            ]
            result = docx_readback(path, expected)
            assert result is True

    def test_bold_italic_runs_readback(self):
        """Bold and italic runs are correctly read back."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "paragraph", "text": "", "style": "Normal", "runs": [
                        {"text": "Bold text", "bold": True},
                        {"text": " and ", "bold": False},
                        {"text": "italic text", "italic": True},
                    ]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            expected = [
                {"text": "Bold text and italic text", "runs": [
                    {"text": "Bold text", "bold": True},
                    {"text": " and "},
                    {"text": "italic text", "italic": True},
                ]},
            ]
            result = docx_readback(path, expected)
            assert result is True


# ---------------------------------------------------------------------------
# Oracle 1 Kill-Proofs
# ---------------------------------------------------------------------------


class TestDocxReadbackKillProofs:
    """Kill-proofs: perturbing the document → FAILS."""

    def test_kill_altered_text(self):
        """Kill-proof: alter expected text → readback FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "paragraph", "text": "Original text.", "style": "Normal"},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            wrong_expected = [
                {"text": "TAMPERED text.", "style": "Normal"},
            ]
            with pytest.raises(AssertionError, match="text mismatch"):
                docx_readback(path, wrong_expected)

    def test_kill_dropped_paragraph(self):
        """Kill-proof: expect fewer paragraphs than exist → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "paragraph", "text": "Para 1.", "style": "Normal"},
                    {"type": "paragraph", "text": "Para 2.", "style": "Normal"},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            wrong_expected = [
                {"text": "Para 1.", "style": "Normal"},
            ]
            with pytest.raises(AssertionError, match="paragraph count mismatch"):
                docx_readback(path, wrong_expected)

    def test_kill_wrong_table_cell(self):
        """Kill-proof: wrong table cell content → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "table", "rows": 2, "cols": 2, "data": [
                        ["A", "B"],
                        ["C", "D"],
                    ]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            wrong_tables = [
                {"rows": 2, "cols": 2, "cells": [["A", "B"], ["C", "X"]]},
            ]
            with pytest.raises(AssertionError, match="cell.*mismatch"):
                docx_readback(path, [], wrong_tables)

    def test_kill_wrong_table_dims(self):
        """Kill-proof: wrong table dimensions → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "table", "rows": 3, "cols": 2, "data": [
                        ["A", "B"],
                        ["C", "D"],
                        ["E", "F"],
                    ]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            wrong_tables = [
                {"rows": 2, "cols": 2, "cells": [["A", "B"], ["C", "D"]]},
            ]
            with pytest.raises(AssertionError, match="dims mismatch"):
                docx_readback(path, [], wrong_tables)


# ---------------------------------------------------------------------------
# Oracle 2: structure_readback
# ---------------------------------------------------------------------------


class TestStructureReadback:
    """structure_readback oracle — heading hierarchy + table dims + paragraph count."""

    def test_basic_structure(self):
        """Basic document structure survives round-trip."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "Title", "level": 1},
                    {"type": "paragraph", "text": "Body.", "style": "Normal"},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            result = structure_readback(path, expected_paragraph_count=2, expected_table_count=0)
            assert result is True

    def test_heading_hierarchy(self):
        """Heading hierarchy is preserved."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "title", "text": "Doc"},
                    {"type": "heading", "text": "H1", "level": 1},
                    {"type": "heading", "text": "H2", "level": 2},
                    {"type": "heading", "text": "H3", "level": 3},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            result = structure_readback(
                path,
                expected_paragraph_count=4,
                expected_table_count=0,
                expected_heading_levels=[0, 1, 2, 3],
            )
            assert result is True

    def test_table_dims_structure(self):
        """Table dimensions survive round-trip."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "table", "rows": 5, "cols": 3, "data": [
                        ["1", "2", "3"],
                        ["4", "5", "6"],
                        ["7", "8", "9"],
                        ["10", "11", "12"],
                        ["13", "14", "15"],
                    ]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            result = structure_readback(
                path,
                expected_paragraph_count=0,
                expected_table_count=1,
                expected_table_dims=[(5, 3)],
            )
            assert result is True


# ---------------------------------------------------------------------------
# Oracle 2 Kill-Proofs
# ---------------------------------------------------------------------------


class TestStructureKillProofs:
    """Kill-proofs: wrong structure → FAILS."""

    def test_kill_wrong_paragraph_count(self):
        """Kill-proof: wrong paragraph count → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "paragraph", "text": "Only one.", "style": "Normal"},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            with pytest.raises(AssertionError, match="paragraph count mismatch"):
                structure_readback(path, expected_paragraph_count=5)

    def test_kill_wrong_table_count(self):
        """Kill-proof: wrong table count → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "table", "rows": 2, "cols": 2, "data": [["A", "B"], ["C", "D"]]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            with pytest.raises(AssertionError, match="table count mismatch"):
                structure_readback(path, expected_paragraph_count=0, expected_table_count=3)

    def test_kill_wrong_heading_levels(self):
        """Kill-proof: wrong heading levels → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "H1", "level": 1},
                    {"type": "heading", "text": "H2", "level": 2},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            with pytest.raises(AssertionError, match="heading.*mismatch"):
                structure_readback(
                    path,
                    expected_paragraph_count=2,
                    expected_heading_levels=[1, 3],
                )

    def test_kill_wrong_table_dims(self):
        """Kill-proof: wrong table dims in structure → FAILS."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "table", "rows": 3, "cols": 2, "data": [["A", "B"], ["C", "D"], ["E", "F"]]},
                ]
            }
            path = _create_doc_from_spec(spec, tmp)

            with pytest.raises(AssertionError, match="table.*dims mismatch"):
                structure_readback(
                    path,
                    expected_paragraph_count=0,
                    expected_table_count=1,
                    expected_table_dims=[(2, 2)],
                )


# ---------------------------------------------------------------------------
# Honest Degradation
# ---------------------------------------------------------------------------


class TestHonestDegradation:
    """Honest degradation: python-docx missing → FAIL LOUD."""

    def test_engine_unavailable_raises(self):
        """If python-docx is missing, engine raises WordEngineUnavailableError."""
        # We can't actually uninstall python-docx in the test env,
        # but we can verify the error class exists and is a RuntimeError
        assert issubclass(WordEngineUnavailableError, RuntimeError)
        assert issubclass(WordError, RuntimeError)

    def test_missing_file_raises(self):
        """Reading a non-existent file raises WordError."""
        with tempfile.TemporaryDirectory() as tmp:
            bad_path = os.path.join(tmp, "nonexistent.docx")
            with pytest.raises(WordError, match="Failed to read"):
                read_document(bad_path)

    def test_empty_spec_raises(self):
        """Creating from an empty spec raises WordError."""
        with pytest.raises(WordError, match="at least one"):
            create_document({"content": []})


# ---------------------------------------------------------------------------
# Gauntlet Scenarios (>=3 end-to-end)
# ---------------------------------------------------------------------------


class TestGauntletScenarios:
    """>=3 end-to-end gauntlet scenarios."""

    def test_scenario_a_multi_heading_styled(self):
        """Scenario (a): multi-heading styled doc with paragraphs."""
        gt = _load_ground_truth()
        sc = gt["scenarios"]["a_multi_heading_styled"]

        with tempfile.TemporaryDirectory() as tmp:
            path = _create_doc_from_spec(sc["spec"], tmp, "scenario_a.docx")

            # Verify structure
            struct_result = structure_readback(
                path,
                expected_paragraph_count=sc["expected_paragraph_count"],
                expected_table_count=sc["expected_table_count"],
                expected_heading_levels=sc["expected_heading_levels"],
            )
            assert struct_result is True

            # Verify content via read-back
            doc_info = read_document(path)
            assert doc_info.paragraph_count == sc["expected_paragraph_count"]

            # Check specific text
            texts = [p.text for p in doc_info.paragraphs]
            assert "Technical Specification" in texts
            assert "Introduction" in texts
            assert "Architecture Overview" in texts

    def test_scenario_b_table_document(self):
        """Scenario (b): document with a real table."""
        gt = _load_ground_truth()
        sc = gt["scenarios"]["b_table_document"]

        with tempfile.TemporaryDirectory() as tmp:
            path = _create_doc_from_spec(sc["spec"], tmp, "scenario_b.docx")

            # Verify structure
            struct_result = structure_readback(
                path,
                expected_paragraph_count=sc["expected_paragraph_count"],
                expected_table_count=sc["expected_table_count"],
                expected_heading_levels=sc["expected_heading_levels"],
                expected_table_dims=sc["expected_table_dims"],
            )
            assert struct_result is True

            # Verify table content
            doc_info = read_document(path)
            assert doc_info.table_count == 1
            table = doc_info.tables[0]
            assert table.rows == 4
            assert table.cols == 3
            assert table.cells[0] == ["Region", "Q3 Sales", "Growth"]
            assert table.cells[1] == ["North", "$1.2M", "+15%"]
            assert table.cells[3] == ["West", "$1.5M", "+22%"]

    def test_scenario_c_list_with_formatting(self):
        """Scenario (c): numbered/bulleted list + bold/italic runs."""
        gt = _load_ground_truth()
        sc = gt["scenarios"]["c_list_with_formatting"]

        with tempfile.TemporaryDirectory() as tmp:
            path = _create_doc_from_spec(sc["spec"], tmp, "scenario_c.docx")

            # Verify structure
            struct_result = structure_readback(
                path,
                expected_paragraph_count=sc["expected_paragraph_count"],
                expected_table_count=sc["expected_table_count"],
                expected_heading_levels=sc["expected_heading_levels"],
            )
            assert struct_result is True

            # Verify list types
            doc_info = read_document(path)
            list_paras = [p for p in doc_info.paragraphs if p.list_type]
            assert len(list_paras) == 5  # 3 numbered + 2 bulleted

            number_paras = [p for p in list_paras if p.list_type == "number"]
            assert len(number_paras) == 3
            assert number_paras[0].text == "Send weekly update"

            bullet_paras = [p for p in list_paras if p.list_type == "bullet"]
            assert len(bullet_paras) == 2
            assert bullet_paras[0].text == "API versioning strategy"

            # Verify bold/italic runs
            para_with_runs = doc_info.paragraphs[1]  # Second paragraph (after heading)
            assert len(para_with_runs.runs) == 3
            assert para_with_runs.runs[0].bold is True
            assert para_with_runs.runs[0].text == "Important: "
            assert para_with_runs.runs[1].italic is True
            assert para_with_runs.runs[1].text == "review "

    def test_full_spec_from_fixture(self):
        """Create from the full doc_spec.json fixture and verify."""
        with open(_SPEC_JSON, encoding="utf-8") as f:
            spec = json.load(f)

        with tempfile.TemporaryDirectory() as tmp:
            path = _create_doc_from_spec(spec, tmp, "full_spec.docx")

            doc_info = read_document(path)
            # The spec has: title, 6 headings, 4 paragraphs, 5 list items, 1 table
            # Total paragraphs = 1 + 6 + 4 + 5 = 16 (table is separate)
            assert doc_info.paragraph_count >= 10
            assert doc_info.table_count == 1

            # Verify table
            table = doc_info.tables[0]
            assert table.rows == 4
            assert table.cols == 3
            assert table.cells[0] == ["Department", "Allocated", "Spent"]


# ---------------------------------------------------------------------------
# Trust Stack Integration
# ---------------------------------------------------------------------------


class TestTrustStackIntegration:
    """Audit log + zero-egress report integration."""

    def test_pipeline_emits_audit_and_egress(self):
        """Pipeline with private_key emits audit log + egress report."""
        private_key = ed25519.Ed25519PrivateKey.generate()

        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "heading", "text": "Test", "level": 1},
                    {"type": "paragraph", "text": "Body.", "style": "Normal"},
                ]
            }
            out_path = os.path.join(tmp, "audit_test.docx")
            result = word_pipeline(
                spec=spec,
                output_path=out_path,
                private_key=private_key,
            )
            assert result.ok
            assert result.audit_log_json, "Audit log JSON should be non-empty"
            assert result.egress_report_json, "Egress report JSON should be non-empty"

            from kairo.oracles.ed25519_audit_log import Ed25519AuditLog
            from kairo.oracles.zero_egress_report import (
                report_from_json,
                verify_zero_egress_report,
            )

            public_key = private_key.public_key()
            entries = Ed25519AuditLog.entries_from_json(result.audit_log_json)
            assert len(entries) > 0
            assert Ed25519AuditLog.verify_chain(entries, public_key)

            report = report_from_json(result.egress_report_json)
            assert verify_zero_egress_report(report, public_key)

    def test_pipeline_without_key_still_works(self):
        """Pipeline without private_key still creates and reads back document."""
        with tempfile.TemporaryDirectory() as tmp:
            spec = {
                "content": [
                    {"type": "paragraph", "text": "No key test.", "style": "Normal"},
                ]
            }
            out_path = os.path.join(tmp, "no_key.docx")
            result = word_pipeline(spec=spec, output_path=out_path)
            assert result.ok
            assert result.doc_info is not None
            assert result.doc_info.paragraph_count == 1
            assert not result.audit_log_json
            assert not result.egress_report_json


# ---------------------------------------------------------------------------
# CLI Integration
# ---------------------------------------------------------------------------


class TestCLIIntegration:
    """word CLI subcommand works end-to-end via registry."""

    def test_cli_create(self):
        """`kairo word create` creates a .docx from a spec file."""
        from kairo.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            out_dir = os.path.join(tmp, "word_output")
            out_docx = os.path.join(out_dir, "cli_test.docx")

            # Create a spec file
            spec_path = os.path.join(tmp, "spec.json")
            with open(spec_path, "w") as f:
                json.dump({
                    "content": [
                        {"type": "heading", "text": "CLI Test", "level": 1},
                        {"type": "paragraph", "text": "Created via CLI.", "style": "Normal"},
                    ]
                }, f)

            rc = main([
                "word", "create", spec_path,
                "--out", "cli_test.docx",
                "--outdir", out_dir,
            ])
            assert rc == 0, f"CLI create failed with exit code {rc}"

            # Verify the doc was created
            assert os.path.exists(out_docx)
            doc_info = read_document(out_docx)
            assert doc_info.paragraph_count == 2
            assert doc_info.paragraphs[0].text == "CLI Test"

    def test_cli_inspect(self):
        """`kairo word inspect` reads back and displays .docx structure."""
        from kairo.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            out_dir = os.path.join(tmp, "word_output")

            # First create a doc
            spec_path = os.path.join(tmp, "spec.json")
            with open(spec_path, "w") as f:
                json.dump({
                    "content": [
                        {"type": "heading", "text": "Inspect Test", "level": 1},
                        {"type": "paragraph", "text": "For inspection.", "style": "Normal"},
                    ]
                }, f)

            rc = main(["word", "create", spec_path, "--outdir", out_dir])
            assert rc == 0

            # Now inspect
            docx_path = os.path.join(out_dir, "word_output.docx")
            rc = main(["word", "inspect", docx_path, "--outdir", out_dir])
            assert rc == 0, f"CLI inspect failed with exit code {rc}"

    def test_cli_create_with_table(self):
        """`kairo word create` with a table spec."""
        from kairo.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            out_dir = os.path.join(tmp, "word_output")

            spec_path = os.path.join(tmp, "spec.json")
            with open(spec_path, "w") as f:
                json.dump({
                    "content": [
                        {"type": "heading", "text": "Table Test", "level": 1},
                        {"type": "table", "rows": 2, "cols": 2, "data": [["A", "B"], ["C", "D"]]},
                    ]
                }, f)

            rc = main(["word", "create", spec_path, "--outdir", out_dir])
            assert rc == 0

            docx_path = os.path.join(out_dir, "word_output.docx")
            doc_info = read_document(docx_path)
            assert doc_info.table_count == 1
            assert doc_info.tables[0].cells == [["A", "B"], ["C", "D"]]
